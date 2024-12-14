#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Sep  5 17:55:35 2024

@author: rafa
"""
import pyttsx3
from bs4 import BeautifulSoup
from selenium import webdriver
from fake_useragent import UserAgent
import codecs
import os
import time
from datetime import datetime, timedelta
from docx import Document
import warnings
import sys
from googletrans import Translator
import docx


def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def createMainFolders(words):
    #CREANDO LA CARPETA DE BUSQUEDA
    searchWords = words
    if not os.path.exists(searchWords):
        os.makedirs(searchWords)

    #CREANDO LA CARPETA JOBS TXT
    pathOutput_txt = words +'/jobs_txt'
    if not os.path.exists(pathOutput_txt):
        os.makedirs(pathOutput_txt)
    
    #CREANDO LA CARPETA JOBS WORD
    pathOutput_docx = words +'/jobs_docx'
    if not os.path.exists(pathOutput_docx):
        os.makedirs(pathOutput_docx)
        
    #CREANDO LA CARPETA DE BUSQUEDA ESPANIOL
    pathOutput_docx_spa = words +'/jobs_doc_spanish'
    if not os.path.exists(pathOutput_docx_spa):
        os.makedirs(pathOutput_docx_spa )

def anyFileExisting(words):
    list_files = os.listdir(words +'/jobs_docx')
    list_files = [x.replace('docx', '') for x in list_files]
    list_files = [x[20:-1] for x in list_files]
    return list_files


def get_current_dictory():
    current_dictory = os.getcwd()
    return current_dictory

def getOfertas(source_text):
    source_text = source_text.split("title")
    for item in source_text:
        if 'description' in item and '[{uid' in item and 'renewedOn' in item:
            item = item.split('[{uid')[0]
            item = item.split('renewedOn')[0]
            MyClass.ofertas.append(item)
            
    with open(words+'/'+'ofertas.txt', 'w') as f:
        for item in MyClass.ofertas:
            f.write('OFERTA')
            f.write('\n')
            f.write(str(item) +'\n')


def getNamesAndLinks(source_text2):
    #SEARCHIN NAMES AND LINKS TOGETER *********************************************************************************
    with open(words+'/'+'Links & Titles.txt', 'w') as f:
        source_text2 = source_text2.split("job-tile-title")
        for item in source_text2:
            if 'link UpLink' in item:
                item = item.split('link UpLink')[1]
                item = item.split('href="')[1]
                link = item.split('">')[0]
                link = 'https://www.upwork.com/nx/search' + link
                title = item.split('">')[1]
                title = title.split('</a>')[0]
                title = title.lstrip()
                f.write(str(link) +'\n')
                f.write(str(title) +'\n')

def getAllData(words, source_text_aux):
    
    for item in MyClass.ofertas:
        
        #TITULO
        item_titulo = item.split(',description')[0]
        item_titulo = item_titulo.replace(':"', "")
        item_titulo = item_titulo.replace('"', "")
        if '<' in item_titulo:
            item_titulo = item_titulo.replace('</span>','')
            item_titulo = item_titulo.replace('<span','')
            item_titulo = item_titulo.replace('class=highlight>','')
            item_titulo = item_titulo.lstrip()

        #DESCRIPTION
        item_description = item.split(',description')[1]
        item_description = item_description.replace(':"', "")
        item_description = item_description.split('",')[0]
        item_description = item_description.replace('<span class="highlight">', "")
        item_description = item_description.replace('</span>', "")
        
        #ITEM CREADO
        item_creado = item.split(',createdOn')[1]
        item_creado = item_creado.replace(':"', "")
        item_creado = item_creado.split('",')[0]

        #ITEM PUBLICADO
        item_publicado= item.split(',publishedOn')[1]
        item_publicado = item_publicado.replace(':"', "")
        item_publicado = item_publicado.split('",')[0]

        
        source_text_links = source_text_aux
        source_text_links = source_text_links.split('up-n-link')
        
        for count, item in enumerate(source_text_links):
            if 'href=' in item and 'referrer_url_path' in item:
                link = item.split('href=')[1]
                link = link.split('>')[0]
                link = link.replace('"','')
                link = 'https://www.upwork.com' + link
        
        # LE SACO EL LINK
        MyClass.data.append([item_titulo, item_description, item_creado, item_publicado])

        try:
            with open(words+'/data/'+item_titulo+'.txt', 'w') as f:          
                f.write(item_titulo +'\n')
                f.write('\n')
                f.write(item_description +'\n')
                f.write('\n')               
                f.write(item_creado +'\n')
                f.write('\n')
                f.write(item_publicado +'\n')
                f.write('\n')
                # f.write(link +'\n')
                # f.write('\n')
        except:
            pass

class MyClass():
    
    ofertas = []
    data = []
    links = []

def main3(words):
    
    try:
        
        warnings.filterwarnings("ignore", category=DeprecationWarning) 
        print('Searching..'+words+' at '+ str(datetime.now()).split('.')[0])
        time.sleep(5)
        words = words.replace(' ', '+')    
        words2 = words.replace('+', ' ')
        createMainFolders(words2)
        list_files_existing = anyFileExisting(words2)
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument("--incognito")
        options.add_argument("--nogpu")
        options.add_argument("--disable-gpu")
        options.add_argument("--window-size=1280,1280")
        options.add_argument("--no-sandbox")
        options.add_argument("--enable-javascript")
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        options.add_argument('--disable-blink-features=AutomationControlled')
        ua = UserAgent()
        userAgent = ua.random
        driver = webdriver.Chrome(options=options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        driver.execute_cdp_cmd('Network.setUserAgentOverride', {"userAgent": userAgent})
        # https://www.upwork.com/nx/search/jobs/?from_recent_search=true&q=cad&sort=recency
        url = 'https://www.upwork.com/nx/search/jobs/?from_recent_search=true&q='+words+'&sort=recency'
        driver.get(url)
        source_page = driver.page_source
        soup = BeautifulSoup(source_page, 'html.parser')
        driver.quit()
        source_text=soup.prettify()
        source_text = codecs.decode(source_text, 'unicode-escape')
        source_text_aux = source_text 
        source_text2 = source_text 
        translator = Translator()
        
        getOfertas(source_text)
        getNamesAndLinks(source_text2)
        
        
        #SOURCE TEXT *********************************************************************************
        with open(words+'/'+'Source.txt', 'w') as f:
            f.write(source_text2)
        
        with open(words+'/'+'Source2.txt', 'w') as f:
            f.write(source_text[3])      
        
        
        
                    

# ***********************************************

        words = words.replace('+', ' ')
        
        getAllData(words, source_text_aux)
        # for item in MyClass.ofertas:

        #     item_titulo = item.split(',description')[0]
        #     item_titulo = item_titulo.replace(':"', "")
        #     item_titulo = item_titulo.replace('"', "")
        #     if '<' in item_titulo:
        #         item_titulo = item_titulo.replace('</span>','')
        #         item_titulo = item_titulo.replace('<span','')
        #         item_titulo = item_titulo.replace('class=highlight>','')
        #         item_titulo = item_titulo.lstrip()
  
        #     titulo_normal = item_titulo
            
            
            # with open(words+'/'+'titulos.txt', 'w') as f:
            #     for item in list_files_existing:
            #         item = item.lstrip()
            #         f.write(str(item) +'\n')
            
            # #WORKING NAMES WEB LINKS *********************************************************************************
            # source_text_links = source_text_aux
            # source_text_links = source_text_links.split('up-n-link')
            
            # for count, item in enumerate(source_text_links):
            #     if 'href=' in item and 'referrer_url_path' in item:
            #         link = item.split('href=')[1]
            #         link = link.split('>')[0]
            #         link = link.replace('"','')
            #         link = 'https://www.upwork.com' + link
            #         MyClass.links.append(link)
    
            # with open(words+'/'+'links.txt', 'w') as f:
            #     for link in MyClass.links:
            #         f.write(link +'\n')
            #         f.write('\n')
            # #END WORKING WEB LINKS *********************************************************************************
            
            

            # for item in MyClass.ofertas:
                
            #     if titulo_normal not in list_files_existing:
                    
            #         item_titulo = item_titulo
                    
            #         item_description = item.split(',description')[1]
            #         item_description = item_description.replace(':"', "")
            #         item_description = item_description.split('",')[0]
            #         item_description = item_description.replace('<span class="highlight">', "")
            #         item_description = item_description.replace('</span>', "")
            #         item_description = "DESCRIPCION: " + item_description
                
            #         item_creado = item.split(',createdOn')[1]
            #         item_creado = item_creado.replace(':"', "")
            #         item_creado = item_creado.split('",')[0]
            #         item_creado = "CREADO: " + item_creado
                    
            #         item_publicado= item.split(',publishedOn')[1]
            #         item_publicado = item_publicado.replace(':"', "")
            #         item_publicado = item_publicado.split('",')[0]
            #         item_publicado = "PUBLICADO: " + item_publicado
            #         MyClass.data.append([item_titulo, item_description, item_creado, item_publicado, link])
                    

                        
                        
                        
                        
            # print('Leng of data:', len(data))
            # print(data[0])
            # print(data[5])
            
            # for count, item in enumerate(MyClass.data):
            #     title = item[0].split('TITULO: ')[1]
            #     title = "".join( x for x in title if (x.isalnum() or x in "._- "))
    
            #     if '/' in title:
            #         title = title.replace('/', "-")
    
            #     date = item[3].split('PUBLICADO: ')[1]
            #     date = date.replace('T', ' ' )
            #     date = date.split('.')[0]
            #     date3 = date
            #     # print(date)
            #     # exit()
            #     start_time = date
            #     replacing_date = datetime.strptime(start_time,'%Y-%m-%d %H:%M:%S')
            #     date = replacing_date+timedelta(hours=-3)
            #     href_txt = MyClass.links[count]
    
                # #CREANDO TXT *********************************************************************************
                # with open(words+'/jobs_txt/'+str(date3)+' - '+title+'.txt', 'w') as f:
                #         f.write(str(item[0]) +'\n')
                #         f.write(str(item[1]) +'\n')
                #         f.write(str(item[2]) +'\n')
                #         f.write(str(item[3]) +'\n')
                #         f.write(str('\n'))
                #         f.write(str(href_txt) +'\n')
    
                # #CREANDO WORD ENGLISH*********************************************************************************
                # nameFile = title
                # description = str(item[1]).replace('DESCRIPCION: ', '' )
                # creado = str(item[2]).replace('CREADO: ', '' )
                # creado = 'CREATED: ' + creado
                # creado = creado.split('.')[0]
                # publicado = str(item[3]).replace('PUBLICADO: ', '' )
                # publicado = 'PUBLISHED: ' + publicado
                # publicado = publicado.split('.')[0]
                # publicado2 = publicado.replace('PUBLISHED: ', '')
                # publicado2 = publicado2.replace(':', '-')    
                # link = href_txt

                # document = Document()
                # document.add_heading(nameFile +'\n', level=1)
                # document.add_paragraph(description +'\n')
                # document.add_paragraph(creado +'\n')
                # document.add_paragraph(publicado +'\n')
                
                # p = document.add_paragraph()
                # add_hyperlink(p, link, nameFile)
               
                # if publicado2 not in list_files_existing:
                #     engine = pyttsx3.init()
                #     engine.say('hello')
                #     # print('One job founded! writing Docx')
                #     document.save(words+'/jobs_docx/' + publicado2 +' '+ nameFile +'.docx')
                
                    
                    
                # #CREANDO WORD SPANISH*********************************************************************************
                # titulo_spa = str(translator.translate(nameFile, dest="es"))
                # titulo_spa = titulo_spa.split('text=')[1]
                # titulo_spa = titulo_spa.split(', pronunciation')[0]
                # titulo_spa = str(titulo_spa)
                # description_spa = str(translator.translate(description, dest="es"))
                # description_spa = description_spa.split('text=')[1]
                # description_spa = description_spa.split(', pronunciation=')[0]
                # creado_esp = 'CREADO: ' + creado
                # creado_esp = creado_esp.replace('CREATED: ', '')
                # publicado_spa = 'PUBLICADO: ' + publicado2
             
                
             
                # document_spa = Document()
                # document_spa.add_heading(titulo_spa +'\n', level=1)
                # document_spa.add_paragraph(description_spa +'\n')
                # document_spa.add_paragraph(creado_esp +'\n')
                # document_spa.add_paragraph(publicado_spa +'\n')
                # p = document_spa.add_paragraph()
                # add_hyperlink(p, link, titulo_spa)
                
                
                
                # if publicado2 not in list_files_existing:
                #     # print('One job founded! writing Docx')
                #     document_spa.save(words+'/jobs_doc_spanish/' + publicado2 +' '+ titulo_spa +'.docx')
    
            # exit()  
              


    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(exc_type, fname, exc_tb.tb_lineno)
        
 
    
def main2(words):
    counterTotal = 0
    while True:
        counterTotal+=1
        print('Excution: ' + str(counterTotal))
        main3(words)
   
def main():

    try:

        print('***Welcome to my Program!***')
        print('***You can now perform a search using keywords. The program will begin the search and will repeat every 5 seconds.***')
        print('***Press ctrl + c to quit***')
        global words
        words = input('Please enter the keywords to search: ')
        main2(words)

    except Exception as e:
        print("Error :", e)
        


if __name__ == '__main__':
    main()


